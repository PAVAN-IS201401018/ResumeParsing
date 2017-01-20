from docx import Document
import copy
import re
from os import listdir
from os.path import isfile, join

def word_to_text(filename):
	document = Document(filename)

#print document.paragraphs
	l=[]
	skills=['python','c','c++','data analyttics','java','programming','machine learning','shell-scripting','advanced java','object oriented','research']
	for p in document.paragraphs:
    		a=p.text

    		b=re.findall(r"[\w']+", a)
    		for i in b:
    			l.append(i)
#print l[0] 
	c=0
	flag=[]
	for i in l:
		if i in skills:
			if i not in flag:
				c=c+1;
				flag.append(i)
	return c


	
def files_from_folder(mypath):
	#mypath="/home/pavan/Desktop/BTP/Test"
	onlyfiles = [f for f in listdir(mypath) if isfile(join(mypath, f))]
	return onlyfiles


k=raw_input("Enter folder path:   ")
l=copy.deepcopy(files_from_folder(k))
print l
for i in l:
	if ".docx" in i:
		print i+" score is  ",
		u=word_to_text(i)
		print u


	
		
		


