from docx import Document
import copy
import re
from os import listdir
from os.path import isfile, join

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from cStringIO import StringIO
skills=['python','c','c++','data analyttics','java','programming','machine learning','shell-scripting','advanced java','object oriented','research']
def convert_pdf_to_txt(mypath):
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = file(mypath, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()

    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
        interpreter.process_page(page)

    text = retstr.getvalue()

    fp.close()
    device.close()
    retstr.close()
    p=text.split('\n')
    m=[]
    for p1 in p:
        if p1:
                a=re.findall(r"[\w']+", p1)
                for i in a:
                        op=i.lower()
                        m.append(op)
    return m






def word_to_text(filename):
	document = Document(filename)

#print document.paragraphs
	l=[]
	for p in document.paragraphs:
    		a=p.text

    		b=re.findall(r"[\w']+", a)
    		for i in b:
			j=i.lower()
    			l.append(j)
	return l



def scorefun(l):
	c=0
	flag=[]
	global skills
	for i in l:
		if i in skills:
			if i not in flag:
				c=c+1;
				flag.append(i)
	return c


	
def files_from_folder(mypath):
	#mypath="/home/pavan/Desktop/BTP/Test"
	onlyfiles = [f for f in listdir(mypath) if isfile(join(mypath, f))]
	return onlyfiles


k=raw_input("Enter folder path:   ")
l=copy.deepcopy(files_from_folder(k))
print l
for i in l:
	if ".docx" in i or ".pdf" in i:
		if ".docx" in i:
			u=word_to_text(i)
		if ".pdf" in i:
			u=convert_pdf_to_txt(i)
		print i+" score is  ",
		s=scorefun(u)
		print s


	
		
		


