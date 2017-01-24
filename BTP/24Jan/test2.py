from docx import Document
import copy
import re
from os import listdir
from os.path import isfile, join

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from cStringIO import StringIO
flag=[]
t=[]
skills=[]
def test(s):
	skills
	if len(skills) == s:
		return 1
def gskillread(j):
	i="/home/pavan/Desktop/BTP/SKILLS/"
	k=str(i)+str(j)
	f=open(k,'r')
	s=[]
	global skills
	skills1=[]
	for line in f:
        	k=line.lower()
        	s.append(k)
	for i in s:
        	k=i.split('\n')
        	if k[0]!='':

               		skills1.append(k[0])
	skills.append(skills1)



 
def skillread():
	global skills
	f=raw_input("Enter the minimum skills required for the candidate:   ")
	g=f.split(' ')
	for i in g:
		j=i.lower()
		skills.append(j)
	#print skills



def convert_pdf_to_txt(i):
    
    j="/home/pavan/Desktop/BTP/Resumes/"
    mypath=str(j)+str(i)
    #  print mypath
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = file(mypath, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()

    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
        interpreter.process_page(page)

    text = retstr.getvalue()

    fp.close()
    device.close()
    retstr.close()
    global t
    p=text.split('\n')
    t=copy.deepcopy(p) 

def afterall():
    m=[]
    global t
    p=copy.deepcopy(t)
    for p1 in p:
        if p1:
                a=re.findall(r"[\w']+", p1)
                for i in a:
                        op=i.lower()
                        m.append(op)
    return m






def word_to_text(filename):
	document = Document(filename)

#print document.paragraphs
	l=[]
	for p in document.paragraphs:
    		a=p.text

    		b=re.findall(r"[\w']+", a)
    		for i in b:
			j=i.lower()
    			l.append(j)
	return l



def scorefun1(a,skills):
	global flag
	l=[]
	s=skills[0]
	skills.pop(0)
	d=s.upper()
	
	for i in a:
		if i in skills and i not in flag:
			j=a.index(i)
			a.pop(j)
			#print i
			flag.append(i)
			l.append(i)
	if len(l)>0:
		print "---------------"+str(d)+"-----------------"
		for i in l:	
			print i
	if len(l)>0:
		print "\n"
	return a

def check():
	global t
	global skills
	for i in skills:
		for j in i:
			for k in t:
				if j in k: 
					hadskills.append(j)
					a=t.index(k)
				
					t.pop(a)

	return hadskills			
	
def files_from_folder(mypath):
	#mypath="/home/pavan/Desktop/BTP/Test"
	onlyfiles = [f for f in listdir(mypath) if isfile(join(mypath, f))]
	return onlyfiles



k="/home/pavan/Desktop/BTP/SKILLS"
l=copy.deepcopy(files_from_folder(k))
#print l
for i in l:
	if ".txt" in i:
		gskillread(i)
a=[]
fi=raw_input("Enter file name ")
print t


if ".docx" in fi:
	word_to_text(fi)
if ".pdf" in fi:
	convert_pdf_to_txt(fi)
print t
hadskills=[]
hadskills=check()
#print hadskills

		
