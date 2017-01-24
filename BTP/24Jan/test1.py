from docx import Document
import copy
import re
from os import listdir
from os.path import isfile, join

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from cStringIO import StringIO
flag=[]
t=[]
#skills=[]
def test(s):
	skills
	if len(skills) == s:
		return 1
def gskillread(k):
	f=open(k,'r')
	s=[]
	skills=[]
	for line in f:
        	k=line.lower()
        	s.append(k)
	for i in s:
        	k=i.split('\n')
        	if k[0]!='':

               		skills.append(k[0])
	return skills



 
def skillread():
	global skills
	f=raw_input("Enter the minimum skills required for the candidate:   ")
	g=f.split(' ')
	for i in g:
		j=i.lower()
		skills.append(j)
	#print skills



def convert_pdf_to_txt(mypath):
    print mypath
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = file(mypath, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()

    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
        interpreter.process_page(page)

    text = retstr.getvalue()

    fp.close()
    device.close()
    retstr.close()
    global t
    p=text.split('\n')
    t=copy.deepcopy(p) 

def afterall():
    m=[]
    global t
    p=copy.deepcopy(t)
    for p1 in p:
        if p1:
                a=re.findall(r"[\w']+", p1)
                for i in a:
                        op=i.lower()
                        m.append(op)
    return m






def word_to_text(filename):
	document = Document(filename)

#print document.paragraphs
	l=[]
	for p in document.paragraphs:
    		a=p.text

    		b=re.findall(r"[\w']+", a)
    		for i in b:
			j=i.lower()
    			l.append(j)
	return l
def scorefun1(a,skills):
	global flag
	l=[]
	s=skills[0]
	skills.pop(0)
	d=s.upper()
	
	for i in a:
		if i in skills and i not in flag:
			j=a.index(i)
			a.pop(j)
			#print i
			flag.append(i)
			l.append(i)
	if len(l)>0:
		print "---------------"+str(d)+"-----------------"
		for i in l:	
			print i
	if len(l)>0:
		print "\n"
	return a
					
	
def files_from_folder(mypath):
	#mypath="/home/pavan/Desktop/BTP/Test"
	onlyfiles = [f for f in listdir(mypath) if isfile(join(mypath, f))]
	return onlyfiles

#skillread()

#k=raw_input("Enter folder path:   ")
k="/home/pavan/Desktop/BTP/Test/23Jan"
l=copy.deepcopy(files_from_folder(k))
#print l
a=[]
fi=raw_input("Enter file name ")
if ".docx" in fi:
	a=word_to_text(fi)
if ".pdf" in fi:
	a=convert_pdf_to_txt(fi)




#print a
n=0
#print l
print "\n"
print "\n"
for i in l:
		if ".txt" in i:
			s=gskillread(i)	
			a=scorefun1(a,s)
			#print i
			n=n+1
#print n			
