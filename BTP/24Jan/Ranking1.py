from docx import Document
import copy
import re
import logging
from os import listdir
from os.path import isfile, join

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from cStringIO import StringIO


skills=[]
def test(s):
	global skills
	if len(skills) == s:
		return 1

 
def skillread():
	global skills
	f=raw_input("Enter the minimum skills required for the candidate:   ")
	g=f.split(' ')
	for i in g:
		j=i.lower()
		skills.append(j)
	print skills



def convert_pdf_to_txt(h):
    j="/home/pavan/Desktop/BTP/Resumes/"
    mypath=str(j)+str(h)
    
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = file(mypath, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()
    try:
    	for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
        	interpreter.process_page(page)
	extractable=1	
    except:  #PDFTextExtractionNotAllowed:
	#print 'This pdf won\'t allow text extraction!'
	extractable=0


    text = retstr.getvalue()

    fp.close()
    device.close()
    retstr.close()
    p=text.split('\n')
    m=[]
    for p1 in p:
        if p1:
                a=re.findall(r"[\w']+", p1)
                for i in a:
                        op=i.lower()
                        m.append(op)
    return m,extractable






def word_to_text(k):
	h="/home/pavan/Desktop/BTP/Resumes/"
	filename=str(h)+str(i)
	
	document = Document(filename)

#print document.paragraphs
	l=[]
	for p in document.paragraphs:
    		a=p.text

    		b=re.findall(r"[\w']+", a)
    		for i in b:
			j=i.lower()
    			l.append(j)
	return l



def scorefun(l):
	c=0
	flag=[]
	global skills
	for i in l:
		if i in skills:
			if i not in flag:
				c=c+1;
				flag.append(i)
	return c


	
def files_from_folder():
	mypath="/home/pavan/Desktop/BTP/Resumes"
	onlyfiles = [f for f in listdir(mypath) if isfile(join(mypath, f))]
	return onlyfiles

skillread()

#k=raw_input("Enter folder path:   ")
l = copy.deepcopy(files_from_folder())
#print l
n=0
for i in l:
	if ".docx" in i or ".pdf" in i:
		if ".docx" in i:
			u=word_to_text(i)
			ex=1
		if ".pdf" in i:
			u,ex=convert_pdf_to_txt(i)
		if ex:
			print i+" score is  ",
			s=scorefun(u)
			n=test(s)
			if n==1:
				print s,
				print " qualified" 
			else:
				print s,
				print  "and not qualified"
		else:
			print i,
			print " is not Extractable"

