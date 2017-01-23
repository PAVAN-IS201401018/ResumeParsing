from docx import Document
import copy
import re
from os import listdir
from os.path import isfile, join

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from cStringIO import StringIO

skills=['ability to work under pressure', 
'accuracy',
'adaptability' 
'administering medication' 
'advising people'
'analyzing data'
'analyzing problems'
'assembling equipment'
'attention to detail'
'auditing financial data' 
'analytical skills'
'attention to details' 
'being thorough'
'brainstorming'
'budgeting'
'building new business'
'business communication skills'
'business management skills'
'calculating data'
'categorizing records'
'checking for accuracy'
'coaching skills'
'collaborating ideas'
'collecting items'
'communicating with young or old people'
'comparing results'
'comprehending books'
'conducting interviews'
'conflict resolution' 
'confronting other people'
'constructing buildings'
'consulting organizations'
'counseling people'
'creative thinking skills'
'creating meaningful work'
'critical thinking skills'
'customer service skills'
'dealing with complaints'
'decision making skills'
'defining problems'
'delegating skills'
'designing systems'
'determination'
'developing plans for projects'
'diplomacy skills'
'displaying art'
'distributing products'
'dramatizing ideas'
'driving safely'
'editing'
'effective listening skills'
'effective study skills'
'encouraging people'
'enforcing rules'
'entertaining others'
'envisioning solutions or ideas'
'estimating project workload'
'ethics'
'evaluating programs'
'expressing feelings'
'expressing ideas'
'extracting information'
'finding missing information'
'following instructions'
'gathering information
'generating accounts
'goal setting 
'initiator 
'handling money
'identifying problems
'imagining innovative solutions
'information management
'inspecting buildings
'inspecting equipment
'interacting with various people 
'interpersonal communication skills 
'interpreting languages
'interviewing 
'inventing products/ideas
'investigating solutions
'knowledge of community 
'knowledge of concepts and principles
'knowledge of government affairs
'leading teams
'listening to people
'maintain focus with interruptions
'maintaining a high level of production
'maintaining accurate records
'maintaining emotional control under stress
'maintaining files
'maintaining schedules or times
'making important decisions
'managing organizations
'managing people 
'mediating between people
'meeting deadlines
'meeting new people 
'motivating others
'multi-tasking
'navigating politics 
'negotiating skills
'operating equipment
'organizing files
'organizing tasks
'patience
'people management skills
'performing clerical work 
'performing numerical analysis
'persuading others
'planning meetings 
'planning organizational needs
'predicting future trends
'preparing written communications
'prioritization skills
'problem analysis skills
'problem solving skills
'product promotion
'promoting events
'proposing ideas
'providing customer service
'providing discipline 
'public speaking
'questioning others
'quick learning skills
'raising funds
'reading 
'recognizing problems
'recruiting
'rehabilitating people
'relating to others'
'reliability'
'remembering information'
'repairing equipment' 
'reporting data'
'researching' 
'resolving conflicts' 
'resourcefulness'
'responsibility'
'results orientated '
'risk taking'
'running meetings'
'sales ability'
'screening telephone calls'
'self-motivated' 
'selling ideas'
'selling products',' services'
'serving people'
'setting performance standards'
'setting up demonstrations'
'sketching charts or diagrams'
'strategic thinking'
'suggesting courses of action'
'summarizing data'
'supervising employees'
'supervising operations'
'supporting others' 
'taking decisive action'
'taking inititiave'
'taking personal responsibility'
'teaching skills' 
'Team building'
'teamwork skills'
'Technical work'
'thinking logically'
'time management'
'training skills'
'translating words'
'using computers'
'verbal communication skills'
'working creatively'  
'working with statistics'
'writing clearly'
'writing letters', 

]
def test(s):
	global skills
	if len(skills) == s:
		return 1

 
def skillread():
	global skills
	f=raw_input("Enter the minimum skills required for the candidate:   ")
	g=f.split(' ')
	for i in g:
		j=i.lower()
		skills.append(j)
	print skills



def convert_pdf_to_txt(mypath):
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = file(mypath, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()

    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
        interpreter.process_page(page)

    text = retstr.getvalue()

    fp.close()
    device.close()
    retstr.close()
    p=text.split('\n')
    m=[]
    for p1 in p:
        if p1:
                a=re.findall(r"[\w']+", p1)
                for i in a:
                        op=i.lower()
                        m.append(op)
    return m






def word_to_text(filename):
	document = Document(filename)

#print document.paragraphs
	l=[]
	for p in document.paragraphs:
    		a=p.text

    		b=re.findall(r"[\w']+", a)
    		for i in b:
			j=i.lower()
    			l.append(j)
	return l



def scorefun(l):
	c=0
	flag=[]
	global skills
	for i in l:
		if i in skills:
			if i not in flag:
				c=c+1;
				flag.append(i)
	return c


	
def files_from_folder(mypath):
	#mypath="/home/pavan/Desktop/BTP/Test"
	onlyfiles = [f for f in listdir(mypath) if isfile(join(mypath, f))]
	return onlyfiles

#skillread()

k=raw_input("Enter folder path:   ")
l=copy.deepcopy(files_from_folder(k))
#print l
n=0
for i in l:
	if ".docx" in i or ".pdf" in i:
		if ".docx" in i:
			u=word_to_text(i)
		if ".pdf" in i:
			u=convert_pdf_to_txt(i)
		print i+" score is  ",
		s=scorefun(u)
		n=test(s)
		if n==1:
			print s,
			print " qualified" 
		else:
			print s,
			print  "and not qualified"

